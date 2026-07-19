# Copyright 2026 Private AI Ltd (SC881079)
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Generated artefacts — DOCX / PDF / MD / HTML / XLSX / PPTX produced for download.
The engines sit behind this swappable interface:

- **DOCX**: pandoc against a styled `reference.docx` (`docx_engine`) is the primary
  route; the structural python-docx builder below is the automatic fallback where
  pandoc is absent. (The reference-document route is the empirical engine choice the
  Decided list anticipated — it produces real Word styles from the model's Markdown.)
- **PDF**: Markdown→HTML→WeasyPrint against a print stylesheet (`pdf_engine`) is the
  primary route; where WeasyPrint's native libraries are absent (e.g. the Windows
  dev box) it falls back to building the DOCX then rendering via LibreOffice.
- **PPTX**: a JSON slide spec built into a 16:9 deck by python-pptx (`pptx_engine`),
  which owns the geometry so the model only chooses content and archetype.

Every non-MD artefact is run through a deterministic validator (`validators`) before
it is returned, so a corrupt file never ships. Both deployment profiles keep working
whatever is installed."""

from __future__ import annotations

import re
from pathlib import Path

from . import render

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Defensive cleanup only (NOT a Markdown→DOCX converter, per): strip the
# emphasis/heading/bullet markers a model may still emit, so they don't render as
# literal "**…**" / "# …" in the document. The Document-drafting skill tells the
# model to emit none; this is the safety net for weaker models.
_EMPHASIS = re.compile(r"\*\*|__|\*|_|`")
_BULLET = re.compile(r"^\s*[-*•]\s+")


def _clean(text: str) -> str:
    return _EMPHASIS.sub("", text).strip()


def _build_docx(title: str, content: str, out_path: str) -> None:
    import docx

    d = docx.Document()
    # Title once: skip the explicit heading when the body already opens with an H1
    # (the block loop renders that `# …` line as the heading instead) — mirrors the
    # docx/pdf engines' `_compose_markdown` dedup so no format double-titles.
    if title and not content.lstrip().startswith("# "):
        d.add_heading(_clean(title), level=0)
    # Blank-line-separated blocks become paragraphs. A block of bullet lines becomes
    # a bullet list; a lone "# heading" line becomes a heading; emphasis markers are
    # stripped. Structural, not a full Markdown parse.
    for block in content.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        lines = [ln for ln in block.splitlines() if ln.strip()]
        if lines and all(_BULLET.match(ln) for ln in lines):
            for ln in lines:
                d.add_paragraph(_clean(_BULLET.sub("", ln)), style="List Bullet")
            continue
        m = re.match(r"^(#{1,6})\s+(.+)$", lines[0]) if lines else None
        if m and len(lines) == 1:
            d.add_heading(_clean(m.group(2)), level=min(len(m.group(1)), 4))
        else:
            d.add_paragraph(_clean(block))
    d.save(out_path)


def generate_artefact(kind: str, title: str, content: str, out_path: str) -> dict:
    """Write an artefact of `kind` (docx|pdf|md|html|xlsx|pptx) to `out_path`; return
    {path, mime}. For pdf, a DOCX is built alongside then converted + removed; for
    html, vendored libs + theme are inlined and a CSP injected (html_engine); for
    xlsx, a JSON workbook spec is built via openpyxl (xlsx_engine); for pptx, a JSON
    slide spec is built via python-pptx (pptx_engine)."""
    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    if kind == "md":
        # Title once: prepend the H1 only when there is a title AND the body does
        # not already open with one (mirrors the docx/pdf `_compose_markdown` rule).
        if title and not content.lstrip().startswith("# "):
            text = f"# {title}\n\n{content}\n"
        else:
            text = f"{content}\n"
        out.write_text(text, encoding="utf-8")
        return {"path": str(out), "mime": "text/markdown"}

    if kind == "html":
        from . import html_engine, validators

        res = html_engine.build(content, title, str(out))
        validators.validate_html(res["path"])
        return res

    if kind == "xlsx":
        from . import validators, xlsx_engine

        res = xlsx_engine.build(content, title, str(out))
        validators.validate_xlsx(res["path"])
        return res

    if kind == "pptx":
        from . import pptx_engine, validators

        res = pptx_engine.build(content, title, str(out))
        validators.validate_pptx(res["path"])
        return res

    if kind == "docx":
        from . import docx_engine, validators

        _write_docx(content, title, str(out))
        validators.validate_docx(str(out))
        return {"path": str(out), "mime": DOCX_MIME}

    if kind == "pdf":
        from . import pdf_engine, validators

        # Primary: WeasyPrint (HTML+CSS Paged Media). Beautiful, browser-free.
        if pdf_engine.available():
            pdf_engine.md_to_pdf(content, title, str(out))
            validators.validate_pdf(str(out))
            return {"path": str(out), "mime": "application/pdf"}

        # Fallback: build the (best-available) DOCX, then render via LibreOffice.
        if not render.available():
            raise RuntimeError("PDF generation needs WeasyPrint or LibreOffice (neither available)")
        docx_tmp = out.with_suffix(".docx")
        _write_docx(content, title, str(docx_tmp))
        # soffice writes <stem>.pdf into the out dir; stem matches → equals out.
        pdf_path = render.docx_to_pdf(str(docx_tmp), str(out.parent))
        try:
            docx_tmp.unlink()
        except OSError:
            pass
        validators.validate_pdf(pdf_path)
        return {"path": pdf_path, "mime": "application/pdf"}

    raise ValueError(f"unknown artefact kind: {kind}")


def _write_docx(content: str, title: str, out_path: str) -> None:
    """Build a DOCX via pandoc+reference.docx where available, else the structural
    python-docx builder. Both produce a valid `.docx` at `out_path`."""
    from . import docx_engine

    if docx_engine.available():
        docx_engine.md_to_docx(content, title, out_path)
    else:
        _build_docx(title, content, out_path)
