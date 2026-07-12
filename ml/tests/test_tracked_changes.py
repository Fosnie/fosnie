"""Correctness tests for the DOCX tracked-change port (lxml). No network.

Builds small DOCX files with python-docx, applies edits, and asserts the
OOXML-level result (valid zip, <w:ins>/<w:del> with author) plus the
accept/reject round-trip and the flatten invariant."""

import zipfile

import docx
import pytest

from app import tracked_changes as tc

W = tracked_changes_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def make_docx(path, paragraphs):
    d = docx.Document()
    for text in paragraphs:
        d.add_paragraph(text)
    d.save(str(path))
    return str(path)


def marker_counts(path):
    """(ins_count, del_count) in word/document.xml."""
    with zipfile.ZipFile(path) as z:
        name = next(n for n in z.namelist() if n.replace("\\", "/") == "word/document.xml")
        xml = z.read(name).decode("utf-8")
    return xml.count(f"<w:ins "), xml.count(f"<w:del ")


def test_apply_creates_tracked_changes(tmp_path):
    src = make_docx(tmp_path / "in.docx", ["The quick brown fox jumps."])
    out = str(tmp_path / "out.docx")
    res = tc.apply_tracked_changes(src, out, [{"find": "quick", "replace": "slow"}], author="Assistant")

    assert len(res["changes"]) == 1
    assert res["errors"] == []
    w_id = res["changes"][0]["w_id"]

    ins, dele = marker_counts(out)
    assert ins == 1 and dele == 1, "one ins + one del expected"

    # Author attribution present.
    with zipfile.ZipFile(out) as z:
        xml = z.read("word/document.xml").decode("utf-8")
    assert 'w:author="Assistant"' in xml
    assert f'w:id="{w_id}"' in xml

    # Re-opens with python-docx (valid DOCX).
    docx.Document(out)

    # Accepted view (read_document) shows the replacement, hides the deletion.
    text = tc.extract_body_text(out)
    assert "slow" in text and "quick" not in text


def test_accept_yields_clean_replacement(tmp_path):
    src = make_docx(tmp_path / "in.docx", ["Pay within 30 days of invoice."])
    proposed = str(tmp_path / "proposed.docx")
    res = tc.apply_tracked_changes(src, proposed, [{"find": "30", "replace": "14"}])
    w_id = res["changes"][0]["w_id"]

    accepted = str(tmp_path / "accepted.docx")
    tc.resolve_tracked_change(proposed, accepted, w_id, "accept")

    ins, dele = marker_counts(accepted)
    assert ins == 0 and dele == 0, "no tracked-change markup after accept"
    text = tc.extract_body_text(accepted)
    assert "14 days" in text and "30 days" not in text


def test_reject_restores_original(tmp_path):
    src = make_docx(tmp_path / "in.docx", ["Pay within 30 days of invoice."])
    proposed = str(tmp_path / "proposed.docx")
    res = tc.apply_tracked_changes(src, proposed, [{"find": "30", "replace": "14"}])
    w_id = res["changes"][0]["w_id"]

    rejected = str(tmp_path / "rejected.docx")
    tc.resolve_tracked_change(proposed, rejected, w_id, "reject")

    ins, dele = marker_counts(rejected)
    assert ins == 0 and dele == 0
    assert tc.extract_body_text(rejected) == "Pay within 30 days of invoice."


def test_accept_all_by_author(tmp_path):
    src = make_docx(tmp_path / "in.docx", ["alpha beta gamma delta"])
    proposed = str(tmp_path / "proposed.docx")
    tc.apply_tracked_changes(
        src,
        proposed,
        [{"find": "alpha", "replace": "ALPHA"}, {"find": "gamma", "replace": "GAMMA"}],
        author="Assistant",
    )
    ins, dele = marker_counts(proposed)
    assert ins == 2 and dele == 2

    out = str(tmp_path / "out.docx")
    res = tc.resolve_all(proposed, out, "accept", author_filter="Assistant")
    assert len(res["resolved"]) == 2
    assert marker_counts(out) == (0, 0)
    text = tc.extract_body_text(out)
    assert "ALPHA" in text and "GAMMA" in text

    # A non-matching author filter resolves nothing.
    out2 = str(tmp_path / "out2.docx")
    res2 = tc.resolve_all(proposed, out2, "accept", author_filter="Nobody")
    assert res2["resolved"] == []
    assert marker_counts(out2) == (2, 2)


def test_pure_deletion(tmp_path):
    src = make_docx(tmp_path / "in.docx", ["Delete this clause entirely please."])
    out = str(tmp_path / "out.docx")
    res = tc.apply_tracked_changes(src, out, [{"find": " entirely", "replace": ""}])
    assert len(res["changes"]) == 1
    ins, dele = marker_counts(out)
    assert ins == 0 and dele == 1
    accepted = str(tmp_path / "acc.docx")
    tc.resolve_tracked_change(out, accepted, res["changes"][0]["w_id"], "accept")
    assert "entirely" not in tc.extract_body_text(accepted)


def test_pure_insertion_with_anchor(tmp_path):
    src = make_docx(tmp_path / "in.docx", ["Section 1. Term."])
    out = str(tmp_path / "out.docx")
    res = tc.apply_tracked_changes(
        src, out, [{"find": "", "replace": " and Termination", "context_before": "Term"}]
    )
    assert len(res["changes"]) == 1
    ins, dele = marker_counts(out)
    assert ins == 1 and dele == 0
    assert "Term and Termination" in tc.extract_body_text(out)


def test_find_not_located_is_reported(tmp_path):
    src = make_docx(tmp_path / "in.docx", ["Nothing to see."])
    out = str(tmp_path / "out.docx")
    res = tc.apply_tracked_changes(src, out, [{"find": "absent", "replace": "x"}])
    assert res["changes"] == []
    assert res["errors"] and res["errors"][0]["index"] == 0


def test_backslash_zip_path_quirk(tmp_path):
    """Older Windows Word stores `word\\document.xml`; the loader must tolerate it."""
    src = make_docx(tmp_path / "normal.docx", ["Backslash quirk test phrase."])
    quirky = tmp_path / "quirky.docx"
    with zipfile.ZipFile(src) as zin, zipfile.ZipFile(quirky, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            data = zin.read(item)
            out_name = "word\\document.xml" if item == "word/document.xml" else item
            zout.writestr(out_name, data)

    out = str(tmp_path / "out.docx")
    res = tc.apply_tracked_changes(str(quirky), out, [{"find": "quirk", "replace": "QUIRK"}])
    assert len(res["changes"]) == 1
    # Re-zipped with normalised forward-slash path.
    with zipfile.ZipFile(out) as z:
        assert "word/document.xml" in z.namelist()
    assert "QUIRK" in tc.extract_body_text(out)


def test_flatten_invariant_matches_plain_text(tmp_path):
    paras = ["First paragraph here.", "Second one follows."]
    src = make_docx(tmp_path / "in.docx", paras)
    assert tc.extract_body_text(src) == "\n".join(paras)


def test_xxe_external_entity_is_not_resolved(tmp_path):
    """A crafted DOCX must NOT be able to read server files via an XML external
    entity (XXE). The hardened parser leaves `&xxe;` unresolved, so the secret
    file's contents can never be inlined into the flattened text."""
    secret = tmp_path / "secret.txt"
    secret.write_text("TOP-SECRET-XXE-CANARY")
    uri = "file:///" + str(secret).replace("\\", "/").lstrip("/")
    payload = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        f'<!DOCTYPE w:document [<!ENTITY xxe SYSTEM "{uri}">]>\n'
        f'<w:document xmlns:w="{W}"><w:body><w:p><w:r><w:t>&xxe;</w:t>'
        "</w:r></w:p></w:body></w:document>"
    ).encode("utf-8")
    mal = tmp_path / "xxe.docx"
    with zipfile.ZipFile(mal, "w") as z:
        z.writestr("word/document.xml", payload)
    try:
        text = tc.extract_body_text(str(mal))
    except Exception:
        return  # parser refused the DTD/entity outright → also safe (fail-closed)
    assert "TOP-SECRET-XXE-CANARY" not in text, "XXE: external entity was resolved"
