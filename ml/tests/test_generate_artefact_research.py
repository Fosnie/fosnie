# Copyright 2026 Private AI Ltd (SC881079)
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Export polish: a corpus report with both
citation namespaces and dual reference sections survives DOCX conversion — the
[D#]/[W#] markers are kept (not stripped as emphasis) and the blank-line-
separated references become one paragraph each, not a single run-on block."""

import pathlib
import sys
import tempfile

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parents[1]))

from docx import Document

from app import generate

_REPORT = """## 1. Findings

The contracts agree on the notice period [D1] while market guidance differs [W1].

## 2. Analysis

Independent reporting corroborates the position [W1], though one document dissents [D2].

## References

### Your documents

[D1] master-agreement.docx — Contracts.

[D2] side-letter.docx — Contracts.

### Web sources

[W1] Market Note 2026 — example.com, published 2026-05-01. https://example.com/a

## Coverage

All 2 documents in scope were reviewed.
"""


def test_md_title_once_when_body_leads_with_h1():
    # The DR finaliser passes content whose single H1 is the title; the MD writer
    # must not prepend a second one.
    body = "# Positioning and Advertising\n\n## 1. Findings\n\nText [W1].\n"
    with tempfile.TemporaryDirectory() as d:
        out = str(pathlib.Path(d) / "r.md")
        generate.generate_artefact("md", "Positioning and Advertising", body, out)
        text = pathlib.Path(out).read_text(encoding="utf-8")
        h1s = [ln for ln in text.splitlines() if ln.startswith("# ")]
        assert h1s == ["# Positioning and Advertising"], f"exactly one H1 title, got {h1s}"


def test_md_title_prepended_when_body_has_none():
    # A plain body (no leading H1) still gets the title prepended once.
    with tempfile.TemporaryDirectory() as d:
        out = str(pathlib.Path(d) / "r.md")
        generate.generate_artefact("md", "My Title", "Just prose, no heading.", out)
        text = pathlib.Path(out).read_text(encoding="utf-8")
        assert [ln for ln in text.splitlines() if ln.startswith("# ")] == ["# My Title"]


def test_docx_fallback_title_once_when_body_leads_with_h1():
    # The python-docx fallback must not add a separate title heading when the body
    # already opens with the title H1.
    body = "# Positioning and Advertising\n\n## 1. Findings\n\nText.\n"
    with tempfile.TemporaryDirectory() as d:
        out = str(pathlib.Path(d) / "r.docx")
        generate._build_docx("Positioning and Advertising", body, out)
        paras = [p.text for p in Document(out).paragraphs]
        assert paras.count("Positioning and Advertising") == 1, f"title once, got {paras}"


def test_docx_keeps_markers_and_per_entry_references():
    with tempfile.TemporaryDirectory() as d:
        out = str(pathlib.Path(d) / "report.docx")
        res = generate.generate_artefact("docx", "Corpus Report", _REPORT, out)
        assert res["path"] == out
        doc = Document(out)
        paras = [p.text for p in doc.paragraphs]
        full = "\n".join(paras)

        # The citation markers survive (brackets are not stripped as emphasis).
        for marker in ("[D1]", "[D2]", "[W1]"):
            assert marker in full, f"{marker} survived DOCX conversion"

        # Both provenance subheadings present.
        assert any("Your documents" in p for p in paras)
        assert any("Web sources" in p for p in paras)

        # Each reference is its OWN paragraph (the blank-line fix), not one run-on.
        ref_paras = [p for p in paras if p.startswith("[D") or p.startswith("[W")]
        assert len(ref_paras) == 3, f"three separate reference paragraphs, got {ref_paras}"
