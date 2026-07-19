"""Beautiful-documents (Phase 1): validators, capability routing/fallback, the
PDF HTML composition, and the zero-egress URL scan over the skill library."""

import re
import zipfile
from pathlib import Path

import docx
import pytest

from app import docx_engine, generate, html_engine, pdf_engine, render, validators

REPO_ROOT = Path(__file__).resolve().parents[2]
SKILLS_DIR = REPO_ROOT / "skills"


# --- Validators --------------------------------------------------------------


def _good_docx(path: Path) -> str:
    d = docx.Document()
    d.add_heading("Title", level=0)
    d.add_paragraph("Body.")
    d.save(str(path))
    return str(path)


def test_validate_docx_accepts_good(tmp_path):
    validators.validate_docx(_good_docx(tmp_path / "ok.docx"))  # no raise


def test_validate_docx_rejects_non_zip(tmp_path):
    p = tmp_path / "bad.docx"
    p.write_bytes(b"this is not a docx")
    with pytest.raises(validators.ValidationError):
        validators.validate_docx(str(p))


def test_validate_docx_rejects_empty(tmp_path):
    p = tmp_path / "empty.docx"
    p.write_bytes(b"")
    with pytest.raises(validators.ValidationError):
        validators.validate_docx(str(p))


def test_validate_docx_rejects_zip_missing_parts(tmp_path):
    p = tmp_path / "wrong.docx"
    with zipfile.ZipFile(p, "w") as z:
        z.writestr("hello.txt", "not a document")
    with pytest.raises(validators.ValidationError):
        validators.validate_docx(str(p))


def test_validate_pdf_rejects_garbage(tmp_path):
    p = tmp_path / "bad.pdf"
    p.write_bytes(b"%NOTPDF")
    with pytest.raises(validators.ValidationError):
        validators.validate_pdf(str(p))


# --- DOCX routing + fallback -------------------------------------------------


def test_docx_falls_back_to_structural_when_pandoc_absent(tmp_path, monkeypatch):
    monkeypatch.setattr(docx_engine, "available", lambda: False)
    out = str(tmp_path / "a.docx")
    res = generate.generate_artefact("docx", "Contract", "Clause one.\n\nClause two.", out)
    assert zipfile.is_zipfile(res["path"])
    body = "\n".join(p.text for p in docx.Document(res["path"]).paragraphs)
    assert "Contract" in body and "Clause one." in body


@pytest.mark.skipif(not docx_engine.available(), reason="pandoc not installed")
def test_docx_pandoc_route_produces_valid_docx(tmp_path):
    out = str(tmp_path / "p.docx")
    md = "## Section\n\nA paragraph.\n\n| A | B |\n| - | - |\n| 1 | 2 |\n"
    res = generate.generate_artefact("docx", "Report", md, out)
    validators.validate_docx(res["path"])  # no raise
    # pandoc rendered the pipe table as a real Word table.
    assert len(docx.Document(res["path"]).tables) >= 1


# --- PDF routing + composition ----------------------------------------------


def test_pdf_errors_when_no_engine(tmp_path, monkeypatch):
    monkeypatch.setattr(pdf_engine, "available", lambda: False)
    monkeypatch.setattr(render, "available", lambda: False)
    with pytest.raises(RuntimeError):
        generate.generate_artefact("pdf", "Brief", "Body.", str(tmp_path / "x.pdf"))


def test_pdf_render_html_inlines_css_and_escapes(tmp_path):
    html = pdf_engine.render_html("# Heading\n\n| A | B |\n| - | - |\n| 1 | 2 |\n", "My <Doc>")
    assert "<style>" in html and "@page" in html          # stylesheet inlined
    assert "<table>" in html                               # gfm table enabled
    assert "&lt;Doc&gt;" in html                           # title escaped
    assert "http://" not in html and "https://" not in html  # zero-egress


@pytest.mark.skipif(not pdf_engine.available(), reason="WeasyPrint native libs absent")
def test_pdf_weasyprint_route(tmp_path):
    out = str(tmp_path / "w.pdf")
    res = generate.generate_artefact("pdf", "Brief", "Body **bold** text.", out)
    assert res["mime"] == "application/pdf"
    with open(res["path"], "rb") as f:
        assert f.read(5) == b"%PDF-"


# --- Zero-egress: no external URLs in the shipped skill library --------------

_URL = re.compile(rb"https?://")


def test_skill_library_has_no_external_urls():
    assert SKILLS_DIR.is_dir(), f"skills library not found at {SKILLS_DIR}"
    offenders = [
        str(f.relative_to(REPO_ROOT))
        for f in SKILLS_DIR.rglob("*")
        if f.is_file() and _URL.search(f.read_bytes())
    ]
    assert not offenders, f"external URL(s) in skill files (zero-egress): {offenders}"


def test_skill_library_ships_phase1_skills():
    slugs = {d.name for d in SKILLS_DIR.iterdir() if (d / "SKILL.md").is_file()}
    assert {"docx-report", "pdf-report", "doc-tables"} <= slugs


# --- HTML artefacts (Phase 2) ------------------------------------------------


def test_html_engine_inlines_echarts_and_injects_csp(tmp_path):
    out = str(tmp_path / "a.html")
    content = (
        "<!DOCTYPE html><html><head><title>T</title>"
        "<!-- pai:theme --><!-- pai:echarts --></head>"
        "<body><div id='c'></div><script>echarts.init(document.getElementById('c'))</script></body></html>"
    )
    res = generate.generate_artefact("html", "T", content, out)
    assert res["mime"] == "text/html"
    doc = Path(res["path"]).read_text(encoding="utf-8")
    assert "<!-- pai:echarts -->" not in doc                      # marker consumed
    assert "Apache Software Foundation" in doc                    # echarts inlined
    assert "Content-Security-Policy" in doc and "default-src 'none'" in doc
    assert "--pai-bg" in doc                                      # theme inlined


def test_html_engine_wraps_a_bare_fragment(tmp_path):
    out = str(tmp_path / "frag.html")
    res = generate.generate_artefact("html", "Frag", "<h1>Hello</h1><p>Body</p>", out)
    doc = Path(res["path"]).read_text(encoding="utf-8")
    assert "<html" in doc.lower() and "<h1>Hello</h1>" in doc
    assert "Content-Security-Policy" in doc


def test_html_engine_echarts_marker_without_vendor_fails(tmp_path, monkeypatch):
    monkeypatch.setattr(html_engine, "_ECHARTS", Path(tmp_path / "missing.js"))
    with pytest.raises(RuntimeError):
        html_engine.build("<html><head><!-- pai:echarts --></head><body></body></html>", "T", str(tmp_path / "x.html"))


def test_validate_html_rejects_external_script(tmp_path):
    p = tmp_path / "cdn.html"
    p.write_text(
        "<html><head><script src='https://cdn.example.com/echarts.js'></script></head><body></body></html>",
        encoding="utf-8",
    )
    with pytest.raises(validators.ValidationError):
        validators.validate_html(str(p))


def test_validate_html_rejects_external_stylesheet(tmp_path):
    p = tmp_path / "css.html"
    p.write_text(
        "<html><head><link rel='stylesheet' href='https://fonts.example/x.css'></head><body></body></html>",
        encoding="utf-8",
    )
    with pytest.raises(validators.ValidationError):
        validators.validate_html(str(p))


def test_validate_html_allows_anchor_links(tmp_path):
    # A hyperlink does not load a resource on render — legitimate report content.
    p = tmp_path / "a.html"
    p.write_text(
        "<html><head></head><body><a href='https://example.com'>source</a></body></html>",
        encoding="utf-8",
    )
    validators.validate_html(str(p))  # no raise


def test_validate_html_rejects_bad_json_island(tmp_path):
    p = tmp_path / "j.html"
    p.write_text(
        "<html><body><script type='application/json'>{not valid json}</script></body></html>",
        encoding="utf-8",
    )
    with pytest.raises(validators.ValidationError):
        validators.validate_html(str(p))


def test_validate_html_accepts_inlined_echarts_page(tmp_path):
    # The real round-trip: an inlined ECharts page (with its namespace URIs and
    # license links) must pass — those are not loading-context URLs.
    out = str(tmp_path / "ok.html")
    content = "<!DOCTYPE html><html><head><!-- pai:echarts --></head><body><div id='c'></div></body></html>"
    res = generate.generate_artefact("html", "Chart", content, out)
    validators.validate_html(res["path"])  # no raise


def test_skill_library_ships_phase2_skills():
    slugs = {d.name for d in SKILLS_DIR.iterdir() if (d / "SKILL.md").is_file()}
    assert {"dashboard", "report-to-page"} <= slugs


def test_echarts_is_vendored():
    assert (REPO_ROOT / "ml" / "app" / "assets" / "vendor" / "echarts.min.js").is_file()


# --- xlsx, ECMA XSD, deterministic pipeline -------------------------


def test_xlsx_round_trip_with_formulas(tmp_path):
    import json

    from openpyxl import load_workbook

    out = str(tmp_path / "wb.xlsx")
    spec = json.dumps({
        "sheets": [{
            "name": "Costs",
            "columns": [{"header": "Item"}, {"header": "Qty"}, {"header": "Unit", "format": "#,##0.00"}, {"header": "Total", "format": "#,##0.00"}],
            "rows": [["Licences", 12, 250, "=B2*C2"], ["Total", "", "", "=SUM(D2:D2)"]],
        }],
    })
    res = generate.generate_artefact("xlsx", "Costs", spec, out)
    assert res["mime"].endswith("spreadsheetml.sheet")
    wb = load_workbook(res["path"])
    ws = wb["Costs"]
    assert ws["A1"].value == "Item"            # header row
    assert ws["D2"].value == "=B2*C2"          # live formula preserved


def test_xlsx_rejects_non_json(tmp_path):
    with pytest.raises(Exception):
        generate.generate_artefact("xlsx", "x", "not json at all", str(tmp_path / "x.xlsx"))


def test_validate_xlsx_rejects_garbage(tmp_path):
    p = tmp_path / "bad.xlsx"
    p.write_bytes(b"not a workbook")
    with pytest.raises(validators.ValidationError):
        validators.validate_xlsx(str(p))


@pytest.mark.skipif(not docx_engine.available(), reason="pandoc not installed")
def test_real_docx_passes_ecma_xsd(tmp_path):
    # Hard assertion (production warns; here we require real pandoc output to validate).
    out = str(tmp_path / "x.docx")
    md = "# Title\n\n## Section\n\nA paragraph with **bold**.\n\n- one\n- two\n\n| A | B |\n| - | - |\n| 1 | 2 |\n"
    generate.generate_artefact("docx", "Report", md, out)
    validators.validate_docx_xsd(out)  # no raise → valid against ISO transitional WML


def test_ooxml_schemas_vendored():
    assert (REPO_ROOT / "ml" / "app" / "assets" / "ooxml-schemas" / "ISO-IEC29500-4_2016" / "wml.xsd").is_file()


def test_generate_pipeline_is_deterministic(tmp_path):
    # The generate→validate pipeline is stable across runs (the no-LLM half of the
    # eval harness). Canned good content; every run must build + validate.
    cases = [
        ("html", "<!DOCTYPE html><html><head><!-- pai:echarts --></head><body><div id='c'></div></body></html>"),
        ("xlsx", '{"sheets":[{"name":"S","columns":[{"header":"A"}],"rows":[[1],["=A2"]]}]}'),
        ("pptx", '{"slides":[{"layout":"title","title":"T"},'
                 '{"layout":"bullets","title":"B","bullets":["one","two"]}]}'),
    ]
    for kind, content in cases:
        for i in range(5):
            out = str(tmp_path / f"{kind}-{i}.{kind}")
            generate.generate_artefact(kind, "Det", content, out)  # raises on any failure
            assert Path(out).is_file()


def test_skill_library_ships_phase3_skills():
    slugs = {d.name for d in SKILLS_DIR.iterdir() if (d / "SKILL.md").is_file()}
    assert {"web-frontend", "research-methods", "xlsx-tables"} <= slugs


def test_skill_library_ships_the_presentations_skill():
    slugs = {d.name for d in SKILLS_DIR.iterdir() if (d / "SKILL.md").is_file()}
    assert {"pptx-deck"} <= slugs
